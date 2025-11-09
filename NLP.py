import os
import re
import json
import fitz  # PyMuPDF
import spacy
from collections import Counter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from xml.sax.saxutils import escape
from docx import Document
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text
import win32com.client as win32  # For DOC to DOCX conversion
import google.generativeai as genai

# ------------------------ CONFIG ------------------------
PDF_FOLDER = "req"
OUTPUT_FOLDER = "outputs"
BATCH_SIZE = 10  # Gemini batch size for passive-to-active
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ------------------------ GEMINI SETUP ------------------------

genai.configure(api_key="AIzaSyD9YVWEXIDux7ugDCCZp8WRta0L0Div-sY")
MODEL_NAME = "gemini-2.5-flash-lite"
model = genai.GenerativeModel(MODEL_NAME)

# ------------------------ LOAD NLP ------------------------
print("Loading spaCy model...")
nlp = spacy.load("en_core_web_sm")


# ------------------------ CONVERT .DOC → .DOCX ------------------------
def convert_doc_to_docx(file_path):
    """Convert old .doc to .docx using Microsoft Word COM interface (Windows only)."""
    try:
        abs_path = os.path.abspath(file_path)
        new_path = abs_path + "x"
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path)
        doc.SaveAs(new_path, FileFormat=16)
        doc.Close()
        word.Quit()
        print(f"🔄 Converted .doc to .docx: {os.path.basename(file_path)}")
        return new_path
    except Exception as e:
        print(f"⚠️ Failed to convert {file_path}: {e}")
        return None


# ------------------------ DETECTORS ------------------------
def is_passive(sent):
    """Detect passive voice including formal ‘shall be’ patterns."""
    text = sent.text.lower()
    passive_patterns = [
        "shall be",
        "is required to be",
        "are provided by",
        "was performed by",
        "were performed by",
        "is defined by",
        "are defined by",
        "is executed by",
    ]
    if any(p in text for p in passive_patterns):
        return True

    passive_aux = any(token.dep_ == "auxpass" for token in sent)
    passive_subj = any(token.dep_ == "nsubjpass" for token in sent)
    vbn_verbs = any(token.tag_ == "VBN" and token.pos_ == "VERB" for token in sent)
    return (passive_aux and passive_subj) or vbn_verbs


def has_conditional_modal(sent):
    modal_words = ["should", "could", "would", "might", "must", "may"]
    return any(token.text.lower() in modal_words for token in sent)


# ------------------------ TEXT EXTRACTION ------------------------
def extract_text_from_pdf(file_path):
    text_pages = []
    pdf = fitz.open(file_path)
    for page_num in range(pdf.page_count):
        page = pdf.load_page(page_num)
        text = " ".join(page.get_text("text").split())
        if text.strip():
            text_pages.append((page_num + 1, text))
    pdf.close()
    return text_pages


def extract_text_from_docx(file_path):
    text = []
    doc = Document(file_path)
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text.strip())
    return [(1, "\n".join(text))] if text else []


def extract_text_from_html(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f, "html.parser")
    for tag in soup(["script", "style"]):
        tag.extract()
    text = soup.get_text(separator=" ")
    text = " ".join(text.split())
    return [(1, text)] if text else []


def extract_text_from_rtf(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        rtf = f.read()
    text = rtf_to_text(rtf)
    text = " ".join(text.split())
    return [(1, text)] if text else []


def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".html", ".htm"]:
        return extract_text_from_html(file_path)
    elif ext == ".rtf":
        return extract_text_from_rtf(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return [(1, text)]
    elif ext == ".doc":
        converted = convert_doc_to_docx(file_path)
        if converted:
            return extract_text_from_docx(converted)
        else:
            return []
    else:
        print(f"⚠️ Unsupported file type: {file_path}")
        return []


# ------------------------ GEMINI: PASSIVE → ACTIVE ------------------------
def passive_to_active(sentences):
    """Convert a batch of passive sentences to active voice using Gemini API."""
    if not sentences:
        return []

    if isinstance(sentences, str):
        sentences = [sentences]

    active_versions = []

    for i in range(0, len(sentences), BATCH_SIZE):
        batch = sentences[i : i + BATCH_SIZE]
        joined = "\n".join([f"{idx + 1}. {s}" for idx, s in enumerate(batch)])
        prompt = f"""
        Convert each of the following sentences from passive to active voice.
        Return ONLY a JSON object with keys as numbers and values as converted sentences.

        Example:
        {{
            "1": "The cat ate the mouse.",
            "2": "The engineer fixed the machine."
        }}

        Sentences:
        {joined}
        """

        for attempt in range(3):
            try:
                response = model.generate_content(prompt)
                text = response.text.strip()
                match = re.search(r"\{[\s\S]*\}", text)
                if not match:
                    raise ValueError("No JSON found in response.")
                json_text = match.group(0)
                json_text = (
                    json_text.replace("\\'", "'")
                    .replace('\\"', '"')
                    .replace("\\n", " ")
                    .replace("\\t", " ")
                )

                try:
                    data = json.loads(json_text)
                except json.JSONDecodeError:
                    print("⚠️ Gemini returned malformed JSON — attempting auto-repair.")
                    repaired = re.sub(r'"\s*"', '", "', json_text)
                    repaired = re.sub(r'(?<=\d)"\s*(?=[A-Za-z])', '", "', repaired)
                    try:
                        data = json.loads(repaired)
                    except Exception:
                        # Fallback manual parsing
                        lines = [
                            line.strip()
                            for line in json_text.split("\n")
                            if ":" in line
                        ]
                        data = {}
                        for line in lines:
                            m = re.match(r'"?(\d+)"?\s*[:\-]\s*"?(.+?)"?[,}]?$', line)
                            if m:
                                data[m.group(1)] = m.group(2).strip()

                for idx in range(1, len(batch) + 1):
                    active_versions.append(data.get(str(idx), "(Conversion missing)"))
                break
            except Exception as e:
                if attempt < 2:
                    print(f"⚠️ Retry Gemini ({attempt+1}/3): {e}")
                else:
                    print(f"❌ Gemini failed: {e}")
                    active_versions.extend([f"(Error: {e})"] * len(batch))
    return active_versions


# ------------------------ CUSTOM SENTENCE SPLITTING ------------------------
def split_sentences_custom(text):
    chunks = re.split(r"(?<=[.;:])\s+|\n+", text)
    return [chunk.strip() for chunk in chunks if len(chunk.strip().split()) >= 3]


# ------------------------ ANALYZER ------------------------
def analyze_file(file_path, file_name):
    results = []
    pages = extract_text(file_path)

    if not pages:
        print(f"⚠️ No readable text found in {file_name}")
        return []

    print(f"📄 Analyzing: {file_name}")

    for page_num, text in pages:
        doc = nlp(text)
        sents = list(doc.sents)
        if len(sents) < 2:
            sents = [nlp(s) for s in split_sentences_custom(text)]

        print(f"   ➤ Sentences found on page {page_num}: {len(sents)}")

        passive_sents = []
        for sent in sents:
            sent_text = sent.text.strip()
            if len(sent_text.split()) < 3:
                continue
            if is_passive(sent):
                passive_sents.append(sent_text)
            if has_conditional_modal(sent):
                results.append(
                    {
                        "File": file_name,
                        "Page": page_num,
                        "Sentence": sent_text,
                        "Issue": "Conditional Modal",
                        "Suggestion": "",
                    }
                )

        if passive_sents:
            actives = passive_to_active(passive_sents)
            for s, a in zip(passive_sents, actives):
                results.append(
                    {
                        "File": file_name,
                        "Page": page_num,
                        "Sentence": s,
                        "Issue": "Passive Voice",
                        "Suggestion": a,
                    }
                )

    return results


# ------------------------ REPORT GENERATION ------------------------
def safe_paragraph(label, text, style):
    return Paragraph(f"<b>{label}</b> {escape(str(text))}", style)


def save_pdf(data):
    pdf_path = os.path.join(OUTPUT_FOLDER, "NLP_Analysis_Report.pdf")
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_path)
    story = []

    story.append(
        Paragraph("<b>Passive Voice & Conditional Modal Detector</b>", styles["Title"])
    )
    story.append(Spacer(1, 20))

    counts = Counter([r["Issue"] for r in data])
    summary_data = [["Issue Type", "Frequency"]] + [[i, c] for i, c in counts.items()]
    table = Table(summary_data, colWidths=[250, 100])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    story.append(Paragraph("<b>Summary of Detected Issues</b>", styles["Heading2"]))
    story.append(table)
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>Detailed Findings</b>", styles["Heading2"]))
    story.append(Spacer(1, 10))

    for r in data:
        story.append(
            safe_paragraph(
                "File & Page:", f"{r['File']} | {r['Page']}", styles["Normal"]
            )
        )
        story.append(safe_paragraph("Sentence:", r["Sentence"], styles["Normal"]))
        story.append(safe_paragraph("Issue:", r["Issue"], styles["Normal"]))
        if r["Suggestion"]:
            story.append(
                safe_paragraph("Active Voice:", r["Suggestion"], styles["Normal"])
            )
        story.append(Spacer(1, 10))

    doc.build(story)
    print(f"✅ Report saved: {pdf_path}")


# ------------------------ MAIN ------------------------
def main():
    print("\n🚀 Starting NLP Text Smell Analysis...\n")
    all_results = []

    for fname in os.listdir(PDF_FOLDER):
        fpath = os.path.join(PDF_FOLDER, fname)
        if fname.lower().endswith(
            (".pdf", ".doc", ".docx", ".html", ".htm", ".rtf", ".txt")
        ):
            try:
                res = analyze_file(fpath, fname)
                all_results.extend(res)
            except Exception as e:
                print(f"⚠️ Error analyzing {fname}: {e}")

    if not all_results:
        print("⚠️ No issues detected or no readable text.")
        return

    save_pdf(all_results)
    print("\n🎯 Analysis complete!")


if __name__ == "__main__":
    main()
